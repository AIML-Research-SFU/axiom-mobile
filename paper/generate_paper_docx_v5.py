#!/usr/bin/env python3
"""
Render paper/PAPER_DRAFT_v5.md into a publication-quality .docx, using the
exact same formatting helpers (Times New Roman, numbered headings, bordered
tables, page numbers) as generate_paper_docx.py used for the v3 paper shown
to the CMPT 416 instructor -- so this looks like the same paper series, not
a different tool's output.

Markdown-driven rather than hardcoded strings (unlike generate_paper_docx.py):
the source of truth is PAPER_DRAFT_v5.md itself, so the .docx can never drift
from the markdown the way hand-porting content between the two risked.

Usage:
    python3 paper/generate_paper_docx_v5.py
"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generate_paper_docx import (
    add_table_with_borders,
    add_heading_numbered,
    add_page_numbers,
)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, "PAPER_DRAFT_v5.md")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "AXIOM_Mobile_Paper_v5.docx")

EM_DASH = "—"
ESC_MARK = "\x01"

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def clean_text(text):
    """Cosmetic cleanup for PLAIN text only -- never call this on text that
    might still contain `code` spans, since -- is a legitimate literal
    character inside code (e.g. `--model`), not a dash to prettify."""
    text = text.replace("\\*", ESC_MARK)
    text = text.replace("--", EM_DASH)
    text = text.replace(ESC_MARK, "*")
    return text


def render_inline_runs(paragraph, text, size=10):
    """Split text on **bold**/`code` spans first, then only apply the
    em-dash/escape cleanup to the plain-text segments -- code spans are
    left byte-for-byte untouched so command-line flags like `--model`
    don't get corrupted into em-dashes before the split ever sees them."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(clean_text(part[2:-2]))
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])  # literal, no clean_text
            run.font.name = "Courier New"
        else:
            run = paragraph.add_run(clean_text(part))
        run.font.size = Pt(size)
        if not run.font.name:
            run.font.name = "Times New Roman"


def add_rich_paragraph(doc, text, size=10, indent=None, list_bullet=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if list_bullet:
        text = "- " + text

    render_inline_runs(p, text, size=size)
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    return p


def parse_table(lines, start):
    header_line = lines[start]
    headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
    i = start + 2
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return headers, rows, i


def strip_inline_md_for_table(text):
    # Strip markdown markers first, then clean -- so a `--flag` inside a
    # table cell doesn't get its dashes prettified into an em dash either.
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = clean_text(text)
    return text


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        raw = f.read()
    lines = raw.split("\n")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(12)

    table_counter = 0
    in_code_block = False
    code_lines = []

    i = 0
    n = len(lines)
    title_done = False
    authors_done = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("# ") and not title_done:
            title_text = stripped[2:].strip()
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_before = Pt(24)
            title_p.paragraph_format.space_after = Pt(8)
            words = title_text.split(" ")
            break_idx = None
            for idx, w in enumerate(words):
                if w.lower() == "for" and idx > 2:
                    break_idx = idx
                    break
            if break_idx:
                line1 = " ".join(words[:break_idx])
                line2 = " ".join(words[break_idx:])
                run = title_p.add_run(line1 + "\n" + line2)
            else:
                run = title_p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            title_done = True
            i += 1
            continue

        if title_done and not authors_done and stripped.startswith("**") and stripped.endswith("**"):
            authors_text = stripped.strip("*")
            author_p = doc.add_paragraph()
            author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_p.paragraph_format.space_after = Pt(2)
            run = author_p.add_run(authors_text)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            i += 1
            while i < n and not lines[i].strip():
                i += 1
            if i < n:
                affil_p = doc.add_paragraph()
                affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                affil_p.paragraph_format.space_after = Pt(16)
                run = affil_p.add_run(lines[i].strip())
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"
                run.italic = True
                i += 1
            authors_done = True
            continue

        if stripped.startswith("*Draft v5"):
            note_p = doc.add_paragraph()
            note_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            note_p.paragraph_format.left_indent = Inches(0.5)
            note_p.paragraph_format.right_indent = Inches(0.5)
            note_p.paragraph_format.space_after = Pt(14)
            run = note_p.add_run(strip_inline_md_for_table(stripped.strip("*")))
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            i += 1
            continue

        if stripped.startswith("## A Note on Why This Draft Exists"):
            add_heading_numbered(doc, "", "Author's Note: Why This Draft Exists")
            doc.paragraphs[-1].runs[0].text = "Author's Note: Why This Draft Exists"
            i += 1
            continue

        if stripped == "## Abstract":
            i += 1
            while i < n and not lines[i].strip():
                i += 1
            abs_text_lines = []
            while i < n and lines[i].strip() and not lines[i].strip().startswith("#") and not lines[i].strip().startswith("---"):
                abs_text_lines.append(lines[i].strip())
                i += 1
            abs_p = doc.add_paragraph()
            abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abs_p.paragraph_format.left_indent = Inches(0.5)
            abs_p.paragraph_format.right_indent = Inches(0.5)
            abs_p.paragraph_format.space_after = Pt(12)
            abs_p.paragraph_format.line_spacing = Pt(12)
            label = abs_p.add_run("Abstract" + EM_DASH)
            label.bold = True
            label.italic = True
            label.font.size = Pt(10)
            label.font.name = "Times New Roman"
            full_text = " ".join(abs_text_lines)
            render_inline_runs(abs_p, full_text, size=10)
            continue

        m = re.match(r"^##\s+(\d+)\.\s+(.+)$", stripped)
        if m:
            add_heading_numbered(doc, m.group(1), m.group(2), level=1)
            i += 1
            continue

        m = re.match(r"^###\s+(\d+\.\d+)\s+(.+)$", stripped)
        if m:
            add_heading_numbered(doc, m.group(1), m.group(2), level=2)
            i += 1
            continue

        m = re.match(r"^##\s+(.+)$", stripped)
        if m and not m.group(1).startswith("A Note"):
            title_txt = m.group(1)
            add_heading_numbered(doc, "", title_txt)
            doc.paragraphs[-1].runs[0].text = title_txt
            i += 1
            continue

        m = re.match(r"^###\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(m.group(1))
            run.bold = True
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            i += 1
            continue

        if stripped.startswith("|"):
            # Look back up to 2 lines for a "**Table: ...**" or
            # "**Table N: ...**" caption line the markdown already has.
            caption_text = None
            for back in range(1, 3):
                if i - back < 0:
                    break
                prev = lines[i - back].strip()
                cm = re.match(r"^\*\*Table[^:]*:\s*(.+?)\*\*\s*$", prev)
                if cm:
                    caption_text = cm.group(1)
                    break
                if prev:
                    break  # non-blank, non-caption line -- stop looking
            headers, rows, i = parse_table(lines, i)
            table_counter += 1
            clean_headers = [strip_inline_md_for_table(h) for h in headers]
            clean_rows = [[strip_inline_md_for_table(c) for c in r] for r in rows]
            add_table_with_borders(doc, clean_headers, clean_rows,
                                    caption=caption_text or "Results, discussed in the surrounding text.",
                                    caption_number=table_counter)
            continue

        if re.match(r"^[-*]\s+", stripped):
            item_text = re.sub(r"^[-*]\s+", "", stripped)
            add_rich_paragraph(doc, item_text, size=10, indent=0.3, list_bullet=True)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Skip a standalone "**Table: ...**" caption line if a table
        # immediately follows -- it's already been used as that table's
        # caption above, so rendering it again here would duplicate it.
        if re.match(r"^\*\*Table[^:]*:.+\*\*$", stripped):
            nxt = i + 1
            while nxt < n and not lines[nxt].strip():
                nxt += 1
            if nxt < n and lines[nxt].strip().startswith("|"):
                i += 1
                continue

        add_rich_paragraph(doc, stripped, size=10)
        i += 1

    add_page_numbers(doc)
    doc.save(OUTPUT_PATH)
    print("Paper saved to: " + OUTPUT_PATH)


if __name__ == "__main__":
    main()
